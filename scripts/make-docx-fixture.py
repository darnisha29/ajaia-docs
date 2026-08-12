#!/usr/bin/env python3
"""Regenerates src/lib/__fixtures__/sample.docx, the .docx import test fixture.

A real Word-format file rather than a hand-written approximation, so the import
tests exercise the same OOXML path a user's upload does.

    pip install python-docx
    python3 scripts/make-docx-fixture.py

Content is chosen to cover every mark and block the editor supports, plus an
escaping case. If you change it, update the assertions in importFile.test.ts.
"""

from pathlib import Path

from docx import Document

OUT = Path(__file__).resolve().parent.parent / "src" / "lib" / "__fixtures__" / "sample.docx"


def main() -> None:
    d = Document()
    d.add_heading("Q3 Product Brief", level=1)
    d.add_heading("Background", level=2)

    p = d.add_paragraph("This paragraph has ")
    p.add_run("bold text").bold = True
    p.add_run(", ")
    p.add_run("italic text").italic = True
    p.add_run(", and ")
    # Underline is the interesting one: mammoth drops it unless style-mapped.
    p.add_run("underlined text").underline = True
    p.add_run(".")

    d.add_heading("Goals", level=3)
    for item in ("Ship the editor", "Validate sharing", "Measure adoption"):
        d.add_paragraph(item, style="List Bullet")

    for item in ("First step", "Second step", "Third step"):
        d.add_paragraph(item, style="List Number")

    d.add_paragraph(
        'A closing paragraph with special chars: <script>alert(1)</script> '
        '& "quotes" and 5 < 10.'
    )
    d.add_paragraph("Tail note.", style="Intense Quote")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
